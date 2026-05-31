"""
Converts eZNR contract Markdown files to beautifully formatted Word documents.
Produces two .docx files: one for Croatian market, one for BiH market.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import re
import os

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)

def add_styled_paragraph(doc, text, style='Normal', bold=False, italic=False, 
                          font_size=None, color=None, alignment=None, space_after=None, space_before=None):
    """Add a styled paragraph to the document."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p

def create_contract_docx(md_path, output_path, country='HR'):
    """Parse markdown and create a formatted Word document."""
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    
    # ── Custom styles ──
    # Heading 1
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)
    
    # Heading 2
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Calibri'
    h2_style.font.size = Pt(13)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x2d, 0x3a, 0x8c)
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(8)
    
    # Heading 3
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Calibri'
    h3_style.font.size = Pt(11.5)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0x37, 0x47, 0x51)
    
    # ── Parse markdown line by line ──
    lines = content.split('\n')
    i = 0
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Horizontal rules (section separators)
        if line.strip() == '---':
            # Add a subtle horizontal line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add border bottom
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '6',
                qn('w:space'): '1',
                qn('w:color'): 'CCCCCC',
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Tables
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            
            # Skip separator rows like |:---|:---|
            if re.match(r'^\|[\s:]*-+', line):
                i += 1
                continue
            
            # Parse cells
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_rows.append(cells)
            
            # Check if next line is NOT a table row
            next_line = lines[i+1].rstrip() if i+1 < len(lines) else ''
            if not next_line.startswith('|'):
                in_table = False
                # Create table
                if table_rows:
                    num_cols = max(len(row) for row in table_rows)
                    tbl = doc.add_table(rows=len(table_rows), cols=num_cols)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tbl.style = 'Table Grid'
                    
                    for r_idx, row in enumerate(table_rows):
                        for c_idx, cell_text in enumerate(row):
                            if c_idx < num_cols:
                                cell = tbl.rows[r_idx].cells[c_idx]
                                # Clean markdown bold
                                clean_text = cell_text.replace('**', '')
                                cell.text = clean_text
                                
                                # Style header row
                                if r_idx == 0:
                                    set_cell_shading(cell, '1a1a2e')
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                            run.font.size = Pt(9.5)
                                            run.font.name = 'Calibri'
                                else:
                                    # Alternate row colors
                                    if r_idx % 2 == 0:
                                        set_cell_shading(cell, 'F5F5FA')
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.size = Pt(9.5)
                                            run.font.name = 'Calibri'
                    
                    # Set column widths proportionally
                    for row in tbl.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.paragraph_format.space_after = Pt(2)
                                paragraph.paragraph_format.space_before = Pt(2)
                    
                    doc.add_paragraph()  # spacing after table
                table_rows = []
            
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        if line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue
        
        if line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue
        
        # Bold label lines (like **1. DAVATELJ LICENCE**)
        bold_match = re.match(r'^\*\*(.+?)\*\*\s*$', line)
        if bold_match:
            p = doc.add_paragraph()
            run = p.add_run(bold_match.group(1))
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            i += 1
            continue
        
        # Lines with mixed bold (like "Tvrtka: ____")
        if '**' in line or '__' in line:
            p = doc.add_paragraph()
            # Parse bold segments
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                else:
                    run = p.add_run(part)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
            i += 1
            continue
        
        # Bullet points
        if line.startswith('- '):
            text = line[2:].strip()
            # Clean markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = text.replace('`', '')
            p = doc.add_paragraph(text, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.name = 'Calibri'
            i += 1
            continue
        
        # Lettered sub-points (a), b), etc.)
        letter_match = re.match(r'^([a-f])\)\s+(.+)', line)
        if letter_match:
            text = f"{letter_match.group(1)}) {letter_match.group(2)}"
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Cm(1.5)
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.name = 'Calibri'
            i += 1
            continue
        
        # Numbered items (1.1., 1.2., etc.)
        num_match = re.match(r'^(\d+\.\d+\.?)\s+(.+)', line)
        if num_match:
            p = doc.add_paragraph()
            num_run = p.add_run(num_match.group(1) + ' ')
            num_run.bold = True
            num_run.font.size = Pt(11)
            num_run.font.name = 'Calibri'
            num_run.font.color.rgb = RGBColor(0x2d, 0x3a, 0x8c)
            
            body_text = num_match.group(2)
            # Handle bold within body text
            parts = re.split(r'(\*\*.*?\*\*)', body_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                else:
                    clean = part.replace('`', '')
                    run = p.add_run(clean)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
            i += 1
            continue
        
        # Italic line (like *sukladno...*)
        italic_match = re.match(r'^\*(.+)\*$', line)
        if italic_match:
            p = doc.add_paragraph()
            run = p.add_run(italic_match.group(1))
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue
        
        # Signature lines (________________________________)
        if '________________________________' in line:
            text = line.replace('________________________________', '_' * 40)
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue
        
        # Regular text
        clean_line = line.strip()
        if clean_line:
            # Remove backticks
            clean_line = clean_line.replace('`', '')
            # Handle inline bold
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', clean_line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                else:
                    run = p.add_run(part)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
        
        i += 1
    
    # ── Save ──
    doc.save(output_path)
    print(f"[OK] Created: {output_path}")


if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    # Croatian contract
    hr_md = os.path.join(base_dir, 'Ugovor_eZNR_Hrvatska.md')
    hr_docx = os.path.join(base_dir, 'Ugovor_eZNR_Hrvatska.docx')
    if os.path.exists(hr_md):
        create_contract_docx(hr_md, hr_docx, country='HR')
    else:
        print(f"[!] Not found: {hr_md}")
    
    # BiH contract
    bih_md = os.path.join(base_dir, 'Ugovor_eZNR_BiH.md')
    bih_docx = os.path.join(base_dir, 'Ugovor_eZNR_BiH.docx')
    if os.path.exists(bih_md):
        create_contract_docx(bih_md, bih_docx, country='BA')
    else:
        print(f"[!] Not found: {bih_md}")
    
    print("\n[DONE] Both contracts generated.")
