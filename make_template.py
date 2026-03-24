import docx
from docx.shared import Pt

def create_template():
    # Load original to keep headers
    doc = docx.Document('C:/Users/zzida/Desktop/znrba/Test ZOP.docx')
    
    # Clear all paragraphs and tables in the document body
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
        
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)
        
    # Add title
    p = doc.add_paragraph('{testTitle}')
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.bold = True
    
    # Add docxtemplater tags
    doc.add_paragraph('{#questions}')
    q_p = doc.add_paragraph('{id}. {text}')
    q_p.runs[0].font.bold = True
    
    doc.add_paragraph('{#options}')
    doc.add_paragraph('{label}) {text}')
    doc.add_paragraph('{/options}')
    doc.add_paragraph('')
    doc.add_paragraph('{/questions}')
    
    doc.add_paragraph('')
    doc.add_paragraph('Potpis kandidata: _________________')
    doc.add_paragraph('Broj tačnih odgovora: ____/____')
    doc.add_paragraph('Kandidat je:')
    doc.add_paragraph('a) Položio')
    doc.add_paragraph('b) Nije položio')
    doc.add_paragraph('Stručno lice koje je vršilo obuku i provjeru znanja:')
    doc.add_paragraph('________________________')
    
    doc.save('C:/Users/zzida/Desktop/znrba/app/public/templates/GeneratedTestTemplate.docx')
    print("Template created!")

create_template()
