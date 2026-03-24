import docx
import json
import re

def extract(file_path):
    print(f"Reading {file_path}")
    doc = docx.Document(file_path)
    questions = []
    
    current_q = None
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt: continue
        
        # Matches '1. QUESTION TEXT?' or '12. QUESTION TEXT?'
        m_q = re.match(r'^(\d+)\.\s*(.*)', txt)
        # Matches 'a) OPTION', 'b) OPTION' 
        m_opt = re.match(r'^([a-k])\)\s*(.*)', txt)
        
        if m_q and len(txt) > 5 and not txt.startswith('1. OBLAST'):
            if current_q: 
                questions.append(current_q)
            current_q = {
                'id': int(m_q.group(1)),
                'text': m_q.group(2).strip(),
                'options': []
            }
        elif m_opt and current_q:
            current_q['options'].append({
                'label': m_opt.group(1),
                'text': m_opt.group(2).strip()
            })
        elif current_q and not m_q and not m_opt:
            if len(current_q['options']) == 0:
                current_q['text'] += ' ' + txt
            else:
                current_q['options'][-1]['text'] += ' ' + txt
                
    if current_q: questions.append(current_q)
    
    # Also parse tables if questions are inside tables (common in word docs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = p.text.strip()
                    if not txt: continue
                    m_q = re.match(r'^(\d+)\.\s*(.*)', txt)
                    m_opt = re.match(r'^([a-k])\)\s*(.*)', txt)
                    
                    if m_q and len(txt) > 5 and not txt.startswith('1. OBLAST'):
                        if current_q: 
                            # Check if we already have it to avoid duplicates
                            if not any(x['id'] == int(m_q.group(1)) for x in questions):
                                questions.append(current_q)
                        current_q = {
                            'id': int(m_q.group(1)),
                            'text': m_q.group(2).strip(),
                            'options': []
                        }
                    elif m_opt and current_q:
                        current_q['options'].append({
                            'label': m_opt.group(1),
                            'text': m_opt.group(2).strip()
                        })
                    elif current_q and not m_q and not m_opt:
                        if len(current_q['options']) == 0:
                            current_q['text'] += ' ' + txt
                        else:
                            current_q['options'][-1]['text'] += ' ' + txt
    if current_q and not any(x['id'] == current_q['id'] for x in questions):
        questions.append(current_q)

    return questions

try:
    znr = extract('c:/Users/zzida/Desktop/znrba/3.1. TEST ZNR.docx')
    zop = extract('c:/Users/zzida/Desktop/znrba/Test ZOP.docx')

    with open('extracted_data.json', 'w', encoding='utf-8') as f:
        json.dump({'znr': znr, 'zop': zop}, f, indent=2, ensure_ascii=False)
    print(f"Extracted {len(znr)} ZNR questions and {len(zop)} ZOP questions")
except Exception as e:
    import traceback
    traceback.print_exc()
