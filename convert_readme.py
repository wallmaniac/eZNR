import os
import re
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# XML helpers for setting background color and cell margins
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_image_safe(doc, path, caption, width, added_images):
    if path in added_images:
        return
    if os.path.exists(path):
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(6)
            p_img.paragraph_format.space_after = Pt(12)
            run_img = p_img.add_run()
            run_img.add_picture(path, width=width)
            
            p_caption = doc.add_paragraph()
            p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_caption.paragraph_format.space_after = Pt(12)
            run_cap = p_caption.add_run(caption)
            run_cap.font.italic = True
            run_cap.font.size = Pt(9.5)
            run_cap.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
            print(f"Successfully added image: {path}")
            added_images.add(path)
        except Exception as e:
            print(f"Warning: Could not add image {path} due to error: {e}")
            # Insert placeholder styled text box
            p_err = doc.add_paragraph()
            p_err.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_err.paragraph_format.space_before = Pt(6)
            p_err.paragraph_format.space_after = Pt(12)
            run_err = p_err.add_run(f"[Prikaz: {caption} - Datoteka: {os.path.basename(path)}]")
            run_err.font.italic = True
            run_err.font.bold = True
            run_err.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

def create_document():
    doc = docx.Document()
    added_images = set()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style Configuration
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Off-black
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # --- COVER PAGE ---
    # Top spacing
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(24)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("eZNR")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(36)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Corporate Blue
    p_title.paragraph_format.space_after = Pt(6)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Sustav za digitalno upravljanje zaštitom na radu (ZNR)")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = RGBColor(0x5C, 0x76, 0x8D)
    p_sub.paragraph_format.space_after = Pt(36)

    # Logo with safe wrapper
    logo_path = "../blulogo.jpg"
    logo_added = False
    try:
        if os.path.exists(logo_path):
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo = p_logo.add_run()
            run_logo.add_picture(logo_path, width=Inches(3.5))
            p_logo.paragraph_format.space_after = Pt(48)
            logo_added = True
    except Exception as e:
        print(f"Warning: Could not add primary logo {logo_path} due to: {e}")

    if not logo_added:
        for alt_logo in ["../extracted_Logo_0.png", "../extracted_Logo_1.png", "../transparentlogo.jpg"]:
            if os.path.exists(alt_logo):
                try:
                    p_logo = doc.add_paragraph()
                    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_logo = p_logo.add_run()
                    run_logo.add_picture(alt_logo, width=Inches(3.5))
                    p_logo.paragraph_format.space_after = Pt(48)
                    logo_added = True
                    print(f"Added alternative logo: {alt_logo}")
                    break
                except Exception as ex:
                    print(f"Warning: Could not add alternative logo {alt_logo} due to: {ex}")

    # Metadata
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(48)
    run_meta = p_meta.add_run("Tehnička specifikacija, arhitektura i poslovni plan\nDatum: Svibanj 2026.\nVerzija: 1.0\nLokacija: Hrvatska / Bosna i Hercegovina")
    run_meta.font.name = 'Calibri'
    run_meta.font.size = Pt(11)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.add_page_break()

    # --- PARSING README.MD ---
    readme_path = "README.md"
    if not os.path.exists(readme_path):
        readme_path = "app/README.md" # fallback
        if not os.path.exists(readme_path):
            print("README.md not found!")
            return

    with open(readme_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_content = []
    table_rows = []
    
    # Custom heading colors & sizes
    h1_color = RGBColor(0x1B, 0x36, 0x5D)
    h2_color = RGBColor(0x2E, 0x5B, 0x88)
    h3_color = RGBColor(0x4A, 0x77, 0x9D)

    def parse_inline_formatting(paragraph, text):
        # Parses basic markdown **bold**, *italic*, and `code` inline formatting
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = paragraph.add_run(part[2:-2])
                r.font.bold = True
            elif part.startswith('*') and part.endswith('*'):
                r = paragraph.add_run(part[1:-1])
                r.font.italic = True
            elif part.startswith('`') and part.endswith('`'):
                r = paragraph.add_run(part[1:-1])
                r.font.name = 'Courier New'
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0xA6, 0x3A, 0x50)
            else:
                paragraph.add_run(part)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\r\n')
        
        # Skip empty lines if they are not in a code block
        if not line.strip() and not in_code_block:
            i += 1
            continue

        # Code block toggle
        if line.startswith("```"):
            if in_code_block:
                # End of code block, render it
                in_code_block = False
                p_code = doc.add_paragraph()
                p_code.paragraph_format.left_indent = Inches(0.5)
                p_code.paragraph_format.right_indent = Inches(0.5)
                p_code.paragraph_format.space_before = Pt(6)
                p_code.paragraph_format.space_after = Pt(6)
                
                # Combine content
                full_code = "\n".join(code_content)
                run_code = p_code.add_run(full_code)
                run_code.font.name = 'Courier New'
                run_code.font.size = Pt(9.5)
                run_code.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                
                # Render code border/bg manually
                p_code_format = p_code.paragraph_format
                p_code_format.line_spacing = 1.0
                
                code_content = []
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Table detection
        if line.startswith("|"):
            table_rows.append(line)
            # Peek next line to see if table continues
            while i + 1 < len(lines) and lines[i+1].startswith("|"):
                i += 1
                table_rows.append(lines[i].rstrip('\r\n'))
            
            # Parse accumulated table rows
            parsed_rows = []
            for tr in table_rows:
                # Strip leading/trailing pipes and split
                cells = [c.strip() for c in tr.strip('|').split('|')]
                parsed_rows.append(cells)
            
            # Filter out separator row (e.g. |:---|:---|)
            filtered_rows = []
            for r_idx, row in enumerate(parsed_rows):
                if r_idx == 1 and all(re.match(r'^:?-+:?$', c) for c in row):
                    continue
                filtered_rows.append(row)
            
            if filtered_rows:
                cols_count = max(len(r) for r in filtered_rows)
                table = doc.add_table(rows=len(filtered_rows), cols=cols_count)
                table.style = 'Table Grid'
                
                for r_idx, row in enumerate(filtered_rows):
                    for c_idx, cell_val in enumerate(row):
                        if c_idx < len(table.rows[r_idx].cells):
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.text = ""
                            p_cell = cell.paragraphs[0]
                            p_cell.paragraph_format.space_after = Pt(2)
                            p_cell.paragraph_format.space_before = Pt(2)
                            
                            # Formatted cell content
                            parse_inline_formatting(p_cell, cell_val)
                            
                            # Styling
                            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                            if r_idx == 0:
                                # Header styling
                                set_cell_background(cell, "1B365D")
                                for run in p_cell.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                    run.font.size = Pt(10)
                            else:
                                # Zebra striping
                                if r_idx % 2 == 0:
                                    set_cell_background(cell, "F2F5F8")
                                    
                doc.add_paragraph() # spacing after table
            table_rows = []
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            # Heading 1 (except the main document title which we handle on cover page)
            heading_text = line[2:]
            if "Sustav za digitalno" not in heading_text:
                p_h = doc.add_paragraph()
                p_h.paragraph_format.space_before = Pt(18)
                p_h.paragraph_format.space_after = Pt(6)
                p_h.paragraph_format.keep_with_next = True
                run_h = p_h.add_run(heading_text)
                run_h.font.name = 'Calibri'
                run_h.font.size = Pt(20)
                run_h.font.bold = True
                run_h.font.color.rgb = h1_color
        elif line.startswith("## "):
            heading_text = line[3:]
            p_h = doc.add_paragraph()
            p_h.paragraph_format.space_before = Pt(14)
            p_h.paragraph_format.space_after = Pt(6)
            p_h.paragraph_format.keep_with_next = True
            run_h = p_h.add_run(heading_text)
            run_h.font.name = 'Calibri'
            run_h.font.size = Pt(15)
            run_h.font.bold = True
            run_h.font.color.rgb = h2_color
            
        elif line.startswith("### "):
            heading_text = line[4:]
            p_h = doc.add_paragraph()
            p_h.paragraph_format.space_before = Pt(10)
            p_h.paragraph_format.space_after = Pt(4)
            p_h.paragraph_format.keep_with_next = True
            run_h = p_h.add_run(heading_text)
            run_h.font.name = 'Calibri'
            run_h.font.size = Pt(12)
            run_h.font.bold = True
            run_h.font.color.rgb = h3_color

        # List items
        elif line.startswith("- ") or line.startswith("* "):
            bullet_text = line[2:]
            p_bullet = doc.add_paragraph(style='List Bullet')
            p_bullet.paragraph_format.space_after = Pt(4)
            parse_inline_formatting(p_bullet, bullet_text)
            
        elif re.match(r'^\d+\.\s', line): # Ordered list item
            num_match = re.match(r'^(\d+\.)\s(.*)', line)
            num_prefix = num_match.group(1)
            bullet_text = num_match.group(2)
            
            p_num = doc.add_paragraph(style='List Number')
            p_num.paragraph_format.space_after = Pt(4)
            parse_inline_formatting(p_num, bullet_text)

        # Standard Paragraph
        else:
            if line.strip():
                p_text = doc.add_paragraph()
                parse_inline_formatting(p_text, line)

        # --- SMART IMAGE HOOK INSERTION SYSTEM ---
        # Scan clean line to inject appropriate images under relevant sections exactly once
        clean_line = line.replace('#', '').replace('*', '').strip()
        
        # 1. Dashboard
        if "Nadzorna ploča" in clean_line or "Dashboard i Pametni Kalendar" in clean_line:
            add_image_safe(doc, "../heropage.png", "Slika 1. Glavna nadzorna ploča (Dashboard) s kalendarom i notifikacijama", Inches(6.0), added_images)
        
        # 2. Zia AI Assistant
        elif "Zia — AI Asistent" in clean_line or "Zia – AI Asistent" in clean_line:
            add_image_safe(doc, "../hzia.png", "Slika 2. Interaktivni chat asistent Zia koji analizira i zadužuje opremu", Inches(6.0), added_images)
        
        # 3. Risk Assessment / Procjena rizika
        elif "Procjena rizika" in clean_line and "Faza" not in clean_line and "Uvođenje" not in clean_line:
            add_image_safe(doc, "../hproc.png", "Slika 3. Modul za Procjenu rizika s interaktivnom 5x5 matricom", Inches(6.0), added_images)
        
        # 4. Certificates / Obuke
        elif "Uvjerenja i certifikati" in clean_line or "Uvjerenja radnika" in clean_line:
            add_image_safe(doc, "../hcert.png", "Slika 4. Izgled generiranog A4 certifikata o osposobljenosti radnika", Inches(6.0), added_images)
        
        # 5. Mobile responsive
        elif "Prilagodba za mobilne" in clean_line or "Mobilna aplikacija i PWA" in clean_line:
            add_image_safe(doc, "../hmob.jpg", "Slika 5. Optimizirano mobilno sučelje (Mobile UI)", Inches(3.2), added_images)

        i += 1

    # Save to the parent folder
    output_filename = "../eZNR_Dokumentacija_i_Poslovni_Plan.docx"
    doc.save(output_filename)
    print(f"Document successfully created: {os.path.abspath(output_filename)}")

if __name__ == "__main__":
    create_document()
